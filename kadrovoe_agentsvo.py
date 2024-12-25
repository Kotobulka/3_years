import sys
import os  
import webbrowser
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QTableWidget, QAbstractItemView,
                             QTableWidgetItem, QPushButton, QFileDialog, 
                             QMessageBox, QDialog, QFormLayout,   
                             QComboBox, QLineEdit, QDateEdit, QMainWindow, QApplication, QTabWidget)  
from PyQt5.QtGui import QIntValidator, QRegularExpressionValidator 
from PyQt5.QtCore import QRegularExpression, QDate, Qt
import psycopg2
from docx import Document
from datetime import datetime 
    

class MainApp(QMainWindow):  
    def __init__(self):  
        super().__init__()  
        self.setWindowTitle("HR App")  
        self.setGeometry(0, 0, 1900, 1060)  

        self.central_widget = QWidget(self)  
        self.setCentralWidget(self.central_widget)  

        self.layout = QVBoxLayout(self.central_widget)  

        # Create tabs  
        self.tabs = QTabWidget(self)  
        self.layout.addWidget(self.tabs)  

        # Create tab instances  
        self.candidates_tab = CandidateTab(self)  
        self.employers_tab = EmployerTab(self)  
        self.vacancies_tab = VacancyTab(self)  
        self.candidate_vacancy_tab = CandidateVacancyTab(self)
        self.report_tab = ReportTab(self)
        '''
        Талица Кандидаты
        Таблица Работодатели
        Таблица Вакансии
        Таблица Предентент_навакансию
        Вкладка отчеты
        '''

        # Add tabs to the interface  
        self.tabs.addTab(self.candidates_tab, "Кандидаты")  
        self.tabs.addTab(self.employers_tab, "Работодатели")  
        self.tabs.addTab(self.vacancies_tab, "Вакансии")  
        self.tabs.addTab(self.candidate_vacancy_tab, "Претенденты на вакансии")
        self.tabs.addTab(self.report_tab, "Отчеты")   

        self.setMinimumSize(800, 600)

class CandidateTab(QWidget):  
    def __init__(self, parent):  
        super().__init__(parent)  
        self.layout = QVBoxLayout(self)  

        self.table = QTableWidget(self)  
        self.table.setColumnCount(11)  # Увеличиваем на 1 для ID  
        self.table.setHorizontalHeaderLabels([  
            "ID", "ФИО", "Дата рождения", "Телефон", "Email", "Резюме",  
            "Опыт работы", "Образование", "Специальность", "Навыки", "Статус"  
        ])  

        self.layout.addWidget(self.table)  
        self.load_data()  

        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)  
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)
        
        # Скрываем столбец ID  
        self.table.hideColumn(0)  

        # Кнопки для управления  
        self.add_button = QPushButton("Добавить")  
        self.add_button.clicked.connect(self.add_candidate)  
        self.layout.addWidget(self.add_button)  

        self.edit_button = QPushButton("Изменить")  
        self.edit_button.clicked.connect(self.edit_candidate)  
        self.layout.addWidget(self.edit_button)  

        self.delete_button = QPushButton("Удалить")  
        self.delete_button.clicked.connect(self.delete_candidate)  
        self.layout.addWidget(self.delete_button)  

        # Кнопка для открытия резюме  
        self.open_resume_button = QPushButton("Открыть резюме")  
        self.open_resume_button.clicked.connect(self.open_resume)  
        self.layout.addWidget(self.open_resume_button)  

    def load_data(self):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute("SELECT id, фио, дата_рождения, телефон, email, резюме, опыт_работы, образование, специальность, навыки, статус FROM Кандидаты ORDER BY id ASC")  
        rows = cursor.fetchall()  

        self.table.setRowCount(len(rows))  
        for i, row in enumerate(rows):  
            for j in range(len(row)):  
                self.table.setItem(i, j, QTableWidgetItem(str(row[j])))  

        cursor.close()  
        connection.close()   

        self.table.resizeColumnsToContents()   

    def add_candidate(self):  
        dialog = CandidateDialog(self, "Добавить кандидата")  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def edit_candidate(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите кандидата для изменения.")  
            return  
        
        # Изменяем диапазон индексов, чтобы включить все 10 полей  
        current_data = [self.table.item(selected_row, i).text() for i in range(1, 11)]  
        candidate_id = self.table.item(selected_row, 0).text()  # ID остается на месте  

        # Отладочный вывод данных  
        print("Редактируемые данные кандидата:", current_data)  

        dialog = CandidateDialog(self, "Изменить кандидата", current_data, candidate_id)  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def delete_candidate(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите кандидата для удаления.")  
            return  

        confirmed = QMessageBox.question(self, "Подтверждение", "Вы уверены, что хотите удалить этого кандидата?", QMessageBox.Yes | QMessageBox.No)  
        if confirmed == QMessageBox.Yes:  
            candidate_id = self.table.item(selected_row, 0).text()  
            self.delete_from_db("Кандидаты", candidate_id)  
            self.load_data()  

    def delete_from_db(self, table_name, candidate_id):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute(f"DELETE FROM {table_name} WHERE id = %s", (candidate_id,))  
        connection.commit()  
        cursor.close()  
        connection.close()  

    def open_resume(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите резюме кандидата для открытия.")  
            return  

        # Получаем путь к резюме  
        resume_path = self.table.item(selected_row, 5).text()  
        
        if not os.path.exists(resume_path):  
            QMessageBox.warning(self, "Ошибка", "Файл резюме по указанному пути не найден. Измените путь или проверьте наличие файла.")  
            return  

        # Открываем резюме в браузере  
        webbrowser.open(resume_path)  

class CandidateDialog(QDialog):  
    def __init__(self, parent, title, candidate_data=None, candidate_id=None):  
        super().__init__(parent)  
        self.setWindowTitle(title)  
        self.layout = QFormLayout(self)  

        self.name_input = QLineEdit(self)  
        self.birthday_input = QDateEdit(self)  
        self.phone_input = QLineEdit(self)  
        self.phone_input.setInputMask("+7(999)999-99-99;_")  # Маска для телефона  
        self.email_input = QLineEdit(self)  
        
        # Устанавливаем валидатор для email  
        email_regex = QRegularExpression(r"^[\w\.-]+@[\w\.-]+\.(ru|com)$")  
        self.email_validator = QRegularExpressionValidator(email_regex, self.email_input)  
        self.email_input.setValidator(self.email_validator)  

        self.resume_input = QLineEdit(self)  
        self.resume_button = QPushButton("Выбрать резюме", self)  
        self.resume_button.clicked.connect(self.select_resume)  # ВЫбор файла (кнопка)  
        
        # Устанавливаем валидатор для опыта работы (только цифры)  
        self.experience_input = QLineEdit(self)  
        self.experience_input.setValidator(QIntValidator(0, 100, self))  # Ввод положительных чисел  

        self.education_input = QComboBox(self)  
        self.education_input.addItems(["Высшее", "СПО"])  # Выпадающий список для образования  
        self.specialization_input = QLineEdit(self)  
        self.skills_input = QLineEdit(self)  
        self.status_input = QComboBox(self)  
        self.status_input.addItems(["Активный", "Архивированный"])  # Выпадающий список для статуса  
        self.candidate_id = candidate_id  

        # Заполняем поля только если переданы данные кандидата  
        if candidate_data and len(candidate_data) >= 10:  
            self.name_input.setText(candidate_data[0])  
            self.birthday_input.setDate(QDate.fromString(candidate_data[1], "yyyy-MM-dd"))   
            self.phone_input.setText(candidate_data[2])  
            self.email_input.setText(candidate_data[3])  
            self.resume_input.setText(candidate_data[4])  
            self.experience_input.setText(candidate_data[5])  
            self.education_input.setCurrentText(candidate_data[6])  
            self.specialization_input.setText(candidate_data[7])  
            self.skills_input.setText(candidate_data[8])  
            self.status_input.setCurrentText(candidate_data[9])  

        self.layout.addRow("ФИО:", self.name_input)  
        self.layout.addRow("Дата рождения:", self.birthday_input)  
        self.layout.addRow("Телефон:", self.phone_input)  
        self.layout.addRow("Email:", self.email_input)  
        self.layout.addRow("Резюме:", self.resume_input)  
        self.layout.addWidget(self.resume_button)
        self.layout.addRow("Опыт работы:", self.experience_input)  
        self.layout.addRow("Образование:", self.education_input)  
        self.layout.addRow("Специальность:", self.specialization_input)  
        self.layout.addRow("Навыки:", self.skills_input)  
        self.layout.addRow("Статус:", self.status_input)  

        self.save_button = QPushButton("Сохранить", self)  
        self.save_button.clicked.connect(self.save)  
        self.layout.addWidget(self.save_button)  

        self.setLayout(self.layout)  

    def select_resume(self):  
        options = QFileDialog.Options()  
        options |= QFileDialog.ReadOnly  
        file_name, _ = QFileDialog.getOpenFileName(self, "Выберите резюме", "", "PDF Files (*.pdf)", options=options)  
        if file_name:  
            self.resume_input.setText(file_name)  # Устанавливаем путь к выбранному резюме в поле ввода  

    def save(self):  
        # Сбор данных из полей  
        name = self.name_input.text()  
        birthday = self.birthday_input.date().toString("yyyy-MM-dd")  
        phone = self.phone_input.text()  
        email = self.email_input.text()  
        resume = self.resume_input.text()  
        experience = self.experience_input.text()   
        education = self.education_input.currentText()  
        specialization = self.specialization_input.text()  
        skills = self.skills_input.text()  
        status = self.status_input.currentText()  

        # Проверка на заполнение всех обязательных полей  
        if not all([name, birthday, phone, email, resume, experience, education, specialization, skills, status]):  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все обязательные поля.")  
            return  

        # Проверка на валидность email  
        if self.email_input.hasAcceptableInput() == False:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, введите корректный Email по шаблону: xxxxxx@xxxx.ru/com.")  
            return  

        # Определяем, добавляем или редактируем  
        is_editing = self.windowTitle() == "Изменить кандидата"  

        # Подключение к базе данных  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  

        if is_editing:  
            # Обновление существующей записи, используя id  
            cursor.execute("""  
                UPDATE Кандидаты   
                SET фио = %s, дата_рождения = %s, телефон = %s, email = %s, резюме = %s,  
                    опыт_работы = %s, образование = %s, специальность = %s, навыки = %s, статус = %s   
                WHERE id = %s  
            """, (name, birthday, phone, email, resume, experience, education, specialization, skills, status, self.candidate_id))  
        else:  
            # Вставка новой записи  
            cursor.execute("""  
                INSERT INTO Кандидаты (фио, дата_рождения, телефон, email, резюме,   
                    опыт_работы, образование, специальность, навыки, статус)   
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)  
            """, (name, birthday, phone, email, resume, experience, education, specialization, skills, status))  

        connection.commit()  
        cursor.close()  
        connection.close()  
        self.accept()  # Закрываем диалог с успешным завершением 

# Employer Tab Class  
class EmployerTab(QWidget):  
    def __init__(self, parent):  
        super().__init__(parent)  
        self.layout = QVBoxLayout(self)  

        self.table = QTableWidget(self)  
        self.table.setColumnCount(4)   
        self.table.setHorizontalHeaderLabels(["ID", "Название", "Контактные данные", "Описание"])  
        self.table.setColumnHidden(0, True)  # Скрываем столбец ID
  
        self.layout.addWidget(self.table)  
        self.load_data()  

        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)

        self.add_button = QPushButton("Добавить")  
        self.add_button.clicked.connect(self.add_employer)  
        self.layout.addWidget(self.add_button)  

        self.edit_button = QPushButton("Изменить")  
        self.edit_button.clicked.connect(self.edit_employer)  
        self.layout.addWidget(self.edit_button)  

        self.delete_button = QPushButton("Удалить")  
        self.delete_button.clicked.connect(self.delete_employer)  
        self.layout.addWidget(self.delete_button)  

    def load_data(self):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute("SELECT id, название, контактные_данные, описание FROM Работодатели ORDER BY id ASC")  
        rows = cursor.fetchall()  

        self.table.setRowCount(len(rows))  
        for i, row in enumerate(rows):  
            for j in range(len(row)):  
                self.table.setItem(i, j, QTableWidgetItem(str(row[j])))  

        cursor.close()  
        connection.close()  

        # Установите сортировку по первой колонке (id)  
        self.table.sortItems(0, Qt.AscendingOrder)
        self.table.resizeColumnsToContents()  

    def add_employer(self):  
        dialog = EmployerDialog(self, "Добавить работодателя")  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def edit_employer(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите работодателя для изменения.")  
            return  
        current_data = [self.table.item(selected_row, i).text() for i in range(4)]  
        dialog = EmployerDialog(self, "Изменить работодателя", current_data)  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def delete_employer(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите работодателя для удаления.")  
            return  

        confirmed = QMessageBox.question(self, "Подтверждение", "Вы уверены, что хотите удалить этого работодателя?", QMessageBox.Yes | QMessageBox.No)  
        if confirmed == QMessageBox.Yes:  
            employer_id = self.table.item(selected_row, 0).text()  # Получаем ID для удаления  
            self.delete_from_db(employer_id)  
            self.load_data()  

    def delete_from_db(self, employer_id):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute("DELETE FROM Работодатели WHERE id = %s", (employer_id,))  
        connection.commit()  
        cursor.close()  
        connection.close()  

class EmployerDialog(QDialog):  
    def __init__(self, parent, title, employer_data=None):  
        super().__init__(parent)  
        self.setWindowTitle(title)  
        self.layout = QFormLayout(self)  

        self.id_input = QLineEdit(self)    
        self.id_input.setVisible(False)    

        self.name_input = QLineEdit(self)  
        self.contact_input = QLineEdit(self)  
        self.description_input = QLineEdit(self)  

        # Устанавливаем валидатор для email  
        email_regex = QRegularExpression(r"^[\w\.-]+@[\w\.-]+\.(ru|com)$")  
        self.email_validator = QRegularExpressionValidator(email_regex, self.contact_input)  
        self.contact_input.setValidator(self.email_validator)  

        if employer_data:  
            self.id_input.setText(employer_data[0])   
            self.name_input.setText(employer_data[1])  
            self.contact_input.setText(employer_data[2])  
            self.description_input.setText(employer_data[3])  
        
        self.layout.addRow("ID:", self.id_input)   
        self.layout.addRow("Название:", self.name_input)  
        self.layout.addRow("Контактные данные:", self.contact_input)  
        self.layout.addRow("Описание:", self.description_input)  

        self.save_button = QPushButton("Сохранить", self)  
        self.save_button.clicked.connect(self.save)  
        self.layout.addWidget(self.save_button)  

        self.setLayout(self.layout)  

    def save(self):  
        employer_id = self.id_input.text()    
        name = self.name_input.text()  
        contact = self.contact_input.text()  
        description = self.description_input.text()  

        # Проверка на заполнение всех обязательных полей  
        if not name or not contact or not description:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все обязательные поля.")  
            return  
        
        # Проверка на маску email  
        if not self.contact_input.hasAcceptableInput():  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, введите корректный Email в формате xxxxxx@xxxxx.ru/com.")  
            return  

        is_editing = self.windowTitle() == "Изменить работодателя"  

        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  

        if is_editing:  
            cursor.execute("""  
                UPDATE Работодатели   
                SET название = %s, контактные_данные = %s, описание = %s   
                WHERE id = %s  
            """, (name, contact, description, employer_id))  
        else:  
            cursor.execute("""  
                INSERT INTO Работодатели (название, контактные_данные, описание)   
                VALUES (%s, %s, %s)  
            """, (name, contact, description))  

        connection.commit()  
        cursor.close()  
        connection.close()  
        self.accept() 

class VacancyTab(QWidget):  
    def __init__(self, parent):  
        super().__init__(parent)  
        self.layout = QVBoxLayout(self)  

        self.table = QTableWidget(self)  
        self.table.setColumnCount(8)  
        self.table.setHorizontalHeaderLabels([  
            "ID", "Название должности", "Описание", "Требования", "Работодатель ID", "Специальность", "Минимальный опыт работы", "Статус"  
        ])  

        self.layout.addWidget(self.table)  
        self.load_data()  
  
        self.table.hideColumn(0)  

        # Устанавливаем режим выбора строк  
        self.table.setSelectionBehavior(QAbstractItemView.SelectRows)  
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)

        # Кнопки для управления  
        self.add_button = QPushButton("Добавить")  
        self.add_button.clicked.connect(self.add_vacancy)  
        self.layout.addWidget(self.add_button)  

        self.edit_button = QPushButton("Изменить")  
        self.edit_button.clicked.connect(self.edit_vacancy)  
        self.layout.addWidget(self.edit_button)  

        self.delete_button = QPushButton("Удалить")  
        self.delete_button.clicked.connect(self.delete_vacancy)  
        self.layout.addWidget(self.delete_button)  

    def load_data(self):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute("SELECT id, название_должности, описание, требования, работодатель_id, специальность, минимальный_опыт, статус FROM Вакансии ORDER BY id ASC")  
        rows = cursor.fetchall()  

        self.table.setRowCount(len(rows))  
        for i, row in enumerate(rows):  
            for j in range(len(row)):  
                self.table.setItem(i, j, QTableWidgetItem(str(row[j])))  

        cursor.close()  
        connection.close()

        self.table.resizeColumnsToContents()  

    def add_vacancy(self):  
        dialog = VacancyDialog(self, "Добавить вакансию")  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def edit_vacancy(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите вакансию для изменения.")  
            return  
        
        vacancy_id = self.table.item(selected_row, 0).text()  
        current_data = [self.table.item(selected_row, i).text() for i in range(1, 8)]    
        dialog = VacancyDialog(self, "Изменить вакансию", current_data, vacancy_id)  
        if dialog.exec_() == QDialog.Accepted:  
            self.load_data()  

    def delete_vacancy(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите вакансию для удаления.")  
            return    

        confirmed = QMessageBox.question(self, "Подтверждение", "Вы уверены, что хотите удалить эту вакансию?", QMessageBox.Yes | QMessageBox.No)  
        if confirmed == QMessageBox.Yes:  
            vacancy_id = self.table.item(selected_row, 0).text()   
            self.delete_from_db("Вакансии", vacancy_id)  
            self.load_data()  

    def delete_from_db(self, table_name, vacancy_id):  
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  
        cursor.execute(f"DELETE FROM {table_name} WHERE id = %s", (vacancy_id,))  
        connection.commit()  
        cursor.close()  
        connection.close()  

class VacancyDialog(QDialog):  
    def __init__(self, parent, title, vacancy_data=None, vacancy_id=None):  
        super().__init__(parent)  
        self.setWindowTitle(title)  
        self.layout = QFormLayout(self)  

        self.vacancy_id = vacancy_id  

        self.title_input = QLineEdit(self)  
        self.description_input = QLineEdit(self)  
        self.requirements_input = QLineEdit(self)  
        self.employer_id_input = QComboBox(self)  
        self.specialty_input = QLineEdit(self)
        self.min_experience_input = QLineEdit(self)    
        self.status_input = QComboBox(self)    

        self.status_input.addItems(["Открыта", "Закрыта"])  

        self.load_employers()  

        if vacancy_data:  
            self.title_input.setText(vacancy_data[0])  
            self.description_input.setText(vacancy_data[1])  
            self.requirements_input.setText(vacancy_data[2])  
            self.employer_id_input.setCurrentText(vacancy_data[3])  
            self.specialty_input.setText(vacancy_data[4])
            self.min_experience_input.setText(str(vacancy_data[5]))  # Заполняем поле минимального опыта работы    
            self.status_input.setCurrentText(vacancy_data[6])    

        self.layout.addRow("Название должности:", self.title_input)  
        self.layout.addRow("Описание:", self.description_input)  
        self.layout.addRow("Требования:", self.requirements_input)  
        self.layout.addRow("Работодатель:", self.employer_id_input)  
        self.layout.addRow("Специальность:", self.specialty_input)
        self.layout.addRow("Минимальный опыт работы:", self.min_experience_input)      
        self.layout.addRow("Статус:", self.status_input)  
  

        self.save_button = QPushButton("Сохранить", self)  
        self.save_button.clicked.connect(self.save)  
        self.layout.addWidget(self.save_button)  

        self.setLayout(self.layout)

    def load_employers(self):  
            connection = psycopg2.connect(  
                dbname='postgres',  
                user='postgres',  
                password='1234',  
                host='localhost',  
                port='5432'  
            )  
            cursor = connection.cursor()  
            cursor.execute("SELECT id, название FROM Работодатели")  
            employers = cursor.fetchall()  

            for employer in employers:  
                self.employer_id_input.addItem(f"{employer[0]} - {employer[1]}")  

            cursor.close()  
            connection.close()  
    def save(self):  
        title = self.title_input.text().strip()  
        description = self.description_input.text().strip()  
        requirements = self.requirements_input.text().strip()  
        employer_id = self.employer_id_input.currentText().split(" - ")[0]  
        specialty = self.specialty_input.text().strip()    
        status = self.status_input.currentText()  

        # Получаем минимальный опыт, проверяем на пустоту  
        min_experience_text = self.min_experience_input.text().strip()  
        min_experience = int(min_experience_text) if min_experience_text else None  # Преобразуем в int или None  

        # Проверка на заполнение всех полей  
        if not title or not description or not requirements or not employer_id or not specialty or not status:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, заполните все поля!")  
            return  

        is_editing = self.windowTitle() == "Изменить вакансию"  
        
        connection = psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  
        cursor = connection.cursor()  

        if is_editing:  
            cursor.execute("""  
                UPDATE Вакансии   
                SET название_должности = %s, описание = %s, требования = %s, специальность = %s,   
                    работодатель_id = %s, статус = %s, минимальный_опыт = %s   
                WHERE id = %s  
            """,   
            (title, description, requirements, specialty, employer_id, status, min_experience, self.vacancy_id))  
        else:  
            cursor.execute("""  
                INSERT INTO Вакансии (название_должности, описание, требования, специальность,   
                    работодатель_id, статус, минимальный_опыт)   
                VALUES (%s, %s, %s, %s, %s, %s, %s)  
            """,  
            (title, description, requirements, specialty, employer_id, status, min_experience))  

        connection.commit()  
        cursor.close()  
        connection.close()  
        self.accept()

class CandidateVacancyTab(QWidget):  
    def __init__(self, parent):  
        super().__init__(parent)  
        self.layout = QVBoxLayout(self)  
 
        self.specialty_combo = QComboBox(self)  
        self.load_specialties()  # Загружаем специальности в комбобокс  
        self.layout.addWidget(self.specialty_combo)  

        self.table = QTableWidget(self)  
        self.table.setColumnCount(3) 
        self.table.setHorizontalHeaderLabels([  
            "ID", "Кандидат ID", "Вакансия ID"  
        ])  

        self.layout.addWidget(self.table)  
        self.load_data()  

        # Скрываем столбец ID  
        self.table.hideColumn(0)  

        # Запрещаем редактирование ячеек таблицы  
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)  

        self.add_button = QPushButton("Добавить")  
        self.add_button.clicked.connect(self.add_candidate_vacancy)  
        self.layout.addWidget(self.add_button)  

        self.edit_button = QPushButton("Изменить")  
        self.edit_button.clicked.connect(self.edit_candidate_vacancy)  
        self.layout.addWidget(self.edit_button)  

        self.delete_button = QPushButton("Удалить")  
        self.delete_button.clicked.connect(self.delete_candidate_vacancy)  
        self.layout.addWidget(self.delete_button)   

    def create_connection(self):  
        """Создание соединения с базой данных."""  
        return psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  

    def load_specialties(self):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        
        cursor.execute("""  
            SELECT DISTINCT специальность   
            FROM (  
                SELECT специальность FROM Кандидаты WHERE специальность IS NOT NULL  
                UNION  
                SELECT специальность FROM Вакансии WHERE специальность IS NOT NULL  
            ) AS combined_specialties  
        """)  
        specialties = cursor.fetchall()  
        
        for specialty in specialties:  
            self.specialty_combo.addItem(specialty[0])  

        cursor.close()  
        connection.close()  

    def load_data(self):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        
        try:  
            cursor.execute("SELECT id, кандидат_id, вакансия_id FROM Претендент_на_вакансию ORDER BY id ASC")  
            rows = cursor.fetchall()  

            self.populate_table(rows)  

        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Произошла ошибка при загрузке данных: {str(e)}")  
        finally:  
            cursor.close()  
            connection.close()    

    def populate_table(self, rows):  
        """Заполнение таблицы данными."""  
        self.table.setRowCount(len(rows))  
        for i, row in enumerate(rows):  
            for j in range(len(row)):  
                self.table.setItem(i, j, QTableWidgetItem(str(row[j])))  

    def add_candidate_vacancy(self):  
        dialog = CandidateVacancyDialog(self)  
        if dialog.exec_() == QDialog.Accepted:  
            data = dialog.get_data()  
            self.insert_into_db(data)  
            self.load_data()  

    def insert_into_db(self, data):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        try:  
            insert_query = """  
            INSERT INTO Претендент_на_вакансию (кандидат_id, вакансия_id, Дата_создания)   
            VALUES (%s, %s, NOW())  
            """  
            cursor.execute(insert_query, data)  
            connection.commit()  
        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Ошибка при добавлении: {str(e)}")  
            connection.rollback()  
        finally:  
            cursor.close()  
            connection.close()  

    def edit_candidate_vacancy(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите запись для изменения.")  
            return  

        id_value = self.table.item(selected_row, 0).text()  
        dialog = CandidateVacancyDialog(self, id_value)  
        if dialog.exec_() == QDialog.Accepted:  
            data = dialog.get_data()  
            self.update_db(id_value, data)  
            self.load_data()  

    def update_db(self, id_value, data):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        try:  
            update_query = """  
            UPDATE Претендент_на_вакансию   
            SET кандидат_id = %s   
            WHERE id = %s  
            """  
            cursor.execute(update_query, (*data, id_value))  
            connection.commit()  
        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Ошибка при обновлении: {str(e)}")  
            connection.rollback()  
        finally:  
            cursor.close()  
            connection.close()  

    def delete_candidate_vacancy(self):  
        selected_row = self.table.currentRow()  
        if selected_row < 0:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите претендента для удаления.")  
            return  

        confirmed = QMessageBox.question(self, "Подтверждение", "Вы уверены, что хотите удалить этого претендента?", QMessageBox.Yes | QMessageBox.No)  
        if confirmed == QMessageBox.Yes:  
            candidate_vacancy_id = self.table.item(selected_row, 0).text()  
            self.delete_from_db(candidate_vacancy_id)  
            self.load_data()  

    def delete_from_db(self, candidate_vacancy_id):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        try:  
            # Получаем id кандидата и вакансии перед удалением  
            cursor.execute("SELECT кандидат_id, вакансия_id FROM Претендент_на_вакансию WHERE id = %s", (candidate_vacancy_id,))  
            candidate_vacancy = cursor.fetchone()  
            if candidate_vacancy:  
                candidate_id, vacancy_id = candidate_vacancy  

                # Обновляем статус кандидата на "Активный"  
                cursor.execute("UPDATE Кандидаты SET статус = 'Активный' WHERE id = %s", (candidate_id,))  
                
                # Обновляем статус вакансии на "Открыта"  
                cursor.execute("UPDATE Вакансии SET статус = 'Открыта' WHERE id = %s", (vacancy_id,))  

                # Удаляем запись из таблицы  
                cursor.execute("DELETE FROM Претендент_на_вакансию WHERE id = %s", (candidate_vacancy_id,))  
                
                connection.commit()  
            else:  
                QMessageBox.warning(self, "Ошибка", "Запись не найдена.")  

        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Ошибка при удалении: {str(e)}")  
            connection.rollback()  
        finally:  
            cursor.close()  
            connection.close()

    def insert_into_db(self, data):  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        try:  
            # Проверка на существование дублирующих записей  
            check_query = """  
                SELECT COUNT(*) FROM Претендент_на_вакансию   
                WHERE кандидат_id = %s AND вакансия_id = %s  
            """  
            cursor.execute(check_query, data)  
            count = cursor.fetchone()[0]  
            
            if count > 0: 
                return  # Прекращаем выполнение, если запись уже существует  

            # Вставка новой записи  
            insert_query = """  
            INSERT INTO Претендент_на_вакансию (кандидат_id, вакансия_id, Дата_создания)   
            VALUES (%s, %s, NOW())  
            """  
            cursor.execute(insert_query, data)  
            connection.commit()  
        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Ошибка при добавлении: {str(e)}")  
            connection.rollback()  
        finally:  
            cursor.close()  
            connection.close()       

class CandidateVacancyDialog(QDialog):  
    def __init__(self, parent, id_value=None):  
        super().__init__(parent)  
        self.id_value = id_value  
        self.setWindowTitle("Добавить Претендента на Вакансию")  
        self.layout = QFormLayout(self)  

        # Комбобокс для выбора вакансии  
        self.vacancy_combo = QComboBox(self)  
        self.vacancy_combo.setEditable(False)  # Запрет ввода данных
        self.vacancy_combo.addItem(None)
        self.load_vacancies()  
        self.vacancy_combo.currentIndexChanged.connect(self.load_candidates)  
        self.layout.addRow("Выберите вакансию:", self.vacancy_combo)  

        # Комбобокс для выбора кандидата  
        self.candidate_combo = QComboBox(self)  
        self.candidate_combo.setEditable(False)  # Запрет ввода данных  
        self.layout.addRow("Выберите кандидата:", self.candidate_combo)  

        self.save_button = QPushButton("Сохранить", self)  
        self.save_button.clicked.connect(self.save)  
        self.layout.addWidget(self.save_button)  

        self.setLayout(self.layout)  

    def load_vacancies(self):  
        """Загрузка вакансий со статусом 'Открыта' в комбобокс."""  
        connection = self.create_connection()  
        cursor = connection.cursor()  
        
        cursor.execute("""  
            SELECT id, название_должности, требования, специальность, минимальный_опыт   
            FROM Вакансии   
            WHERE статус = 'Открыта'  
        """)  
        vacancies = cursor.fetchall()  

        if not vacancies:  
            self.vacancy_combo.addItem("Нет доступных вакансий", None)  # Уведомление о недоступности вакансий  
        else:  
            for vacancy in vacancies:  
                display_text = f"{vacancy[0]} - {vacancy[1]}, {vacancy[2]}, {vacancy[3]}"  
                self.vacancy_combo.addItem(display_text, vacancy[0])  # Сохраняем ID вакансии как пользовательские данные  

        cursor.close()  
        connection.close()  

    def load_candidates(self):  
        """Загрузка кандидатов на основе выбранной вакансии."""  
        self.candidate_combo.clear()  # Очищаем предыдущие данные  
        vacancy_id = self.vacancy_combo.currentData()  # Получаем ID выбранной вакансии  

        if vacancy_id is None:  
            return  # Если вакансия не выбрана, выходим  

        connection = self.create_connection()  
        cursor = connection.cursor()  

        # Получаем минимальный опыт и специальность вакансии  
        cursor.execute("SELECT минимальный_опыт, специальность FROM Вакансии WHERE id = %s", (vacancy_id,))  
        vacancy_data = cursor.fetchone()  

        if vacancy_data:  
            min_experience, specialty = vacancy_data  
            cursor.execute("""  
                SELECT id, фио, опыт_работы   
                FROM Кандидаты   
                WHERE статус = 'Активный' AND опыт_работы >= %s AND специальность = %s  
            """, (min_experience, specialty))  
            candidates = cursor.fetchall()  

            if not candidates:  
                self.candidate_combo.addItem("Нет доступных кандидатов", None)  # Уведомление о недоступности кандидатов  
            else:  
                for candidate in candidates:  
                    self.candidate_combo.addItem(f"{candidate[0]} - {candidate[1]} (Опыт: {candidate[2]})", candidate[0])  

        else:  
            self.candidate_combo.addItem("Нет информации о вакансии", None)  # Уведомление о недоступной вакансии  

        cursor.close()  
        connection.close()  

    def save(self):  
        """Сохранение записи о кандидате на вакансию и обновление статусов."""  
        vacancy_id = self.vacancy_combo.currentData()  
        candidate_id = self.candidate_combo.currentData()  

        if vacancy_id is None or candidate_id is None:  
            QMessageBox.warning(self, "Ошибка", "Пожалуйста, выберите вакансию и кандидата!")  
            return  

        connection = self.create_connection()  
        cursor = connection.cursor()  

        try:  
            # Проверка на существование дублирующей записи  
            cursor.execute("""  
                SELECT COUNT(*) FROM Претендент_на_вакансию   
                WHERE кандидат_id = %s AND вакансия_id = %s  
            """, (candidate_id, vacancy_id))  
            count = cursor.fetchone()[0]  
            
            if count > 0:  
                QMessageBox.warning(self, "Ошибка", "Эта запись уже существует.")  
                return  # Прекращаем выполнение, если запись уже существует  

            # Добавление записи о кандидате на вакансию  
            cursor.execute("""  
                INSERT INTO Претендент_на_вакансию (кандидат_id, вакансия_id, Дата_создания)   
                VALUES (%s, %s, NOW())  
            """, (candidate_id, vacancy_id))  

            # Обновление статуса вакансии на "Закрыта"  
            cursor.execute("""  
                UPDATE Вакансии   
                SET статус = 'Закрыта'   
                WHERE id = %s  
            """, (vacancy_id,))  

            # Обновление статуса кандидата на "Архивированный"  
            cursor.execute("""  
                UPDATE Кандидаты   
                SET статус = 'Архивированный'   
                WHERE id = %s  
            """, (candidate_id,))  

            connection.commit()  
            QMessageBox.information(self, "Успех", "Запись успешно добавлена!")  
            self.accept()  # Закрываем диалог  

        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f"Ошибка при добавлении: {str(e)}")  
            connection.rollback()  
        finally:  
            cursor.close()  
            connection.close()  

    def get_data(self):  
        """Возвращает данные о кандидате и вакансии."""  
        vacancy_id = self.vacancy_combo.currentData()  
        candidate_id = self.candidate_combo.currentData()  
        return (candidate_id, vacancy_id)  

    def create_connection(self):  
        """Создание соединения с базой данных."""  
        return psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234'  
        )

class ReportTab(QWidget):  
    def __init__(self, parent):  
        super().__init__(parent)  
        self.layout = QVBoxLayout(self)  

        self.create_report_button = QPushButton("Создать новый отчет", self)  
        self.create_report_button.clicked.connect(self.open_report_dialog)  
        self.layout.addWidget(self.create_report_button)  

        self.open_reports_folder_button = QPushButton("Открыть папку с отчетами", self)  
        self.open_reports_folder_button.clicked.connect(self.open_reports_folder)  
        self.layout.addWidget(self.open_reports_folder_button)  

    def open_report_dialog(self):  
        dialog = ReportDialog(self)  
        dialog.exec_()  # Открываем диалог как модальное окно  

    def open_reports_folder(self):  
        reports_folder = r"C:\Users\Vo1\Downloads\Отчеты"  # Убедитесь, что путь верный  
        if os.path.exists(reports_folder):  
            webbrowser.open(reports_folder)  # Открываем проводник к папке  
        else:  
            QMessageBox.critical(self, "Ошибка", "Папка отчетов не найдена.")

class ReportDialog(QDialog):  
    def __init__(self, parent=None):  
        super().__init__(parent)  
        self.setWindowTitle("Создание отчета")  
        self.setFixedSize(400, 300)  

        self.layout = QFormLayout(self)  

        # Выбор таблицы  
        self.table_combo = QComboBox(self)  
        self.table_combo.addItems(["Кандидаты", "Работодатели", "Вакансии", "Претенденты на вакансии"])  
        self.table_combo.currentIndexChanged.connect(self.update_status_combo)  # Обновляем статусы при изменении таблицы  
        self.layout.addRow("Выберите таблицу:", self.table_combo)  

        # Выбор статуса  
        self.status_combo = QComboBox(self)  
        self.update_status_combo()  # Инициализируем статусы  
        self.layout.addRow("Статус:", self.status_combo)  

        # Выбор начальной и конечной даты  
        self.start_date_edit = QDateEdit(self)  
        self.start_date_edit.setCalendarPopup(True)  
        self.start_date_edit.setDate(QDate.currentDate())  
        self.layout.addRow("Начальная дата:", self.start_date_edit)  

        self.end_date_edit = QDateEdit(self)  
        self.end_date_edit.setCalendarPopup(True)  
        self.end_date_edit.setDate(QDate.currentDate())  
        self.layout.addRow("Конечная дата:", self.end_date_edit)  

        # Кнопки  
        self.generate_button = QPushButton("Сгенерировать отчет", self)  
        self.generate_button.clicked.connect(self.generate_report)  
        self.layout.addRow(self.generate_button)  

        self.cancel_button = QPushButton("Отмена", self)  
        self.cancel_button.clicked.connect(self.reject)  
        self.layout.addRow(self.cancel_button)  

        self.setLayout(self.layout)  

    def update_status_combo(self):  
        table_name = self.table_combo.currentText()  

        # Сброс статусов  
        self.status_combo.clear()  

        # Заполняем статус в зависимости от выбранной таблицы  
        if table_name == "Кандидаты":  
            self.status_combo.addItems(["Все", "Активный", "Архивированный"])  
        elif table_name == "Работодатели":  
            self.status_combo.addItems(["Все"])  
        elif table_name == "Вакансии":  
            self.status_combo.addItems(["Все", "Открыта", "Закрыта"])  
        elif table_name == "Претенденты на вакансии":  
            self.status_combo.addItems(["Все"])  

    def generate_report(self):  
        table_name = self.table_combo.currentText()  
        start_date = self.start_date_edit.date().toString("yyyy-MM-dd")  
        end_date = self.end_date_edit.date().toString("yyyy-MM-dd")  
        status = self.status_combo.currentText()  
        
        try:  
            conn = self.connect_to_db()  # Открываем соединение с БД  
            query = self.create_query(table_name, start_date, end_date, status)  
            records, headers = self.execute_query(conn, query)  
            
            # Сохраняем отчет с датами  
            report_path = self.create_docx_report(table_name, records, headers, start_date, end_date)  

            QMessageBox.information(self, "Успех", f'Отчет сгенерирован: {report_path}')  
            self.accept()  # Закрываем диалог  
        except Exception as e:  
            QMessageBox.critical(self, "Ошибка", f'Ошибка: {str(e)}')  
        finally:  
            conn.close()  # Закрываем соединение  

    def connect_to_db(self):  
        return psycopg2.connect(  
            dbname='postgres',  
            user='postgres',  
            password='1234',  
            host='localhost',  
            port='5432'  
        )  

    def create_query(self, table_name, start_date, end_date, status):  
        query = f"SELECT * FROM {table_name} WHERE дата_создания BETWEEN '{start_date}' AND '{end_date}'"  
        if table_name in ['Кандидаты', 'Вакансии'] and status != "Все":  
            query += f" AND статус = '{status}'"  
        return query  

    def execute_query(self, conn, query):  
        cursor = conn.cursor()  
        cursor.execute(query)  
        records = cursor.fetchall()  
        column_names = [desc[0] for desc in cursor.description]  
        cursor.close()  
        return records, column_names  

    def create_docx_report(self, table_name, records, headers, start_date, end_date):  
        output_dir = r"C:\Users\Vo1\Downloads\Отчеты"  # Указываем путь для сохранения  
        if not os.path.exists(output_dir):  
            os.makedirs(output_dir)  

        # Добавляем даты в название отчета  
        report_path = os.path.join(output_dir, f'{table_name}_report_{start_date}_to_{end_date}.docx')  
        
        # Создаем документ  
        doc = Document()  
        doc.add_heading(f'Отчет по таблице {table_name} с {start_date} по {end_date}', level=1)  

        # Форматируем каждую запись  
        formatted_records = []  
        for record in records:  
            formatted_record = []  
            for item in record:  
                if isinstance(item, datetime):  # Проверяет, является ли элемент объектом datetime  
                    formatted_item = item.strftime("%Y-%m-%d")  # Форматирует дату  
                elif isinstance(item, str):  
                    formatted_item = item  # Оставляет строку без изменений  
                else:  
                    formatted_item = str(item)  # Приводит другие типы к строке  
                
                formatted_record.append(formatted_item)  

            formatted_records.append(formatted_record)  

        if formatted_records:  # Проверяем, есть ли записи  
            doc.add_paragraph(", ".join(headers))  # Добавляем заголовки в документ  
            for record in formatted_records:  
                doc.add_paragraph(", ".join(record))  # Записываем отформатированную запись  

        doc.save(report_path)  # Сохраняем файл  
        return report_path  

if __name__ == "__main__":  
    app = QApplication(sys.argv)  
    main_app = MainApp()  
    main_app.show()  
    sys.exit(app.exec_()) 